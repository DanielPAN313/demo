from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "24h全天交互排期Agent_黑客松方案.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width


def add_fixed_table(doc, rows, headers=None, widths=None, header_fill="E8EEF5"):
    table = doc.add_table(rows=1 if headers else 0, cols=len(headers or rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        hdr = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr[i].text = text
            set_cell_shading(hdr[i], header_fill)
            for p in hdr[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.bold = True
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9.5)
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
            for p in row[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_kv_table(doc, items):
    rows = [[k, v] for k, v in items]
    return add_fixed_table(
        doc,
        rows,
        headers=["项目", "建议"],
        widths=[Inches(1.45), Inches(5.0)],
        header_fill="EEF2F7",
    )


def add_callout(doc, title, body, fill="F6F8FB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="C9D6E2")
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    run2.font.name = "Microsoft YaHei"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run2.font.size = Pt(10)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(46, 116, 181)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for r in runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(10.5)
    return p


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run("24h 全天交互 + 排期模块")
    tr.bold = True
    tr.font.name = "Microsoft YaHei"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    tr.font.size = Pt(24)
    tr.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("Agent Builders Hackathon 南京站 · 赛道二创意智能体方案与明日开发执行稿")
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(85, 85, 85)

    add_callout(
        doc,
        "一句话定位",
        "我们不是做一个更会提醒人的 AI 日历，而是做一个知道什么时候不该打扰你的 24 小时个人节律 Agent。",
        fill="EFF6FF",
    )

    add_h1(doc, "1. 赛道贴合度判断")
    add_para(
        doc,
        "结论：贴合赛道二，而且比赛道一更适合。赛道二强调跳出行业框架、从个人兴趣和奇思出发，做任意天马行空的 AI 智能体。这个项目如果被表达成“AI 日历”，会显得普通；如果表达成“守护注意力和个人节律的 24 小时智能体”，就更符合创意赛道。",
        bold_prefix="结论：",
    )
    add_bullets(
        doc,
        [
            "核心差异不是“自动排任务”，而是“判断该不该打扰用户”。",
            "系统同时管理时间、状态、任务优先级、外部事件和触达边界。",
            "评委容易记住的概念是：注意力防火墙、状态感知排期、夜间静默晨报。",
        ]
    )

    add_h1(doc, "2. 明日黑客松开发判断")
    add_kv_table(
        doc,
        [
            ("现场时间", "2026 年 6 月 28 日 10:00-19:00，实际可开发约 6-8 小时。"),
            ("三人可完成", "强演示型 MVP：前端体验 + Agent 决策逻辑 + 模拟数据 + 路演剧本。"),
            ("不要尝试", "真实穿戴设备接入、完整账户系统、全量 GitHub 监听、多平台真实通知推送。"),
            ("产品级周期", "如果做成真实可长期使用的产品，至少需要 3-6 周。"),
            ("明日目标", "不是功能全，而是故事闭环、界面完整、重排瞬间有冲击力。"),
        ]
    )

    add_h1(doc, "3. MVP 功能范围")
    add_h2(doc, "必须完成")
    add_bullets(
        doc,
        [
            "今日节律首页：用一天时间轴展示专注、碎片、休息、睡眠四类时段。",
            "状态感知重排：输入“睡眠差 / 压力高 / 紧急任务插入”，自动生成新计划。",
            "注意力防火墙：把通知分成立即打扰、稍后提醒、晨报汇总、静默归档。",
            "可解释排期：每次调整都给出理由，例如“睡眠质量低，上午深度任务减少 30%”。",
            "Agent 对话入口：用户说“我现在很累，帮我重排下午”，系统返回新排期和解释。",
        ]
    )
    add_h2(doc, "可以模拟，不必真实接入")
    add_bullets(
        doc,
        [
            "睡眠评分、压力等级、疲劳状态可以用按钮或滑块模拟。",
            "GitHub 紧急 Issue 可以用一条预置事件模拟。",
            "课表、任务列表、协作消息可以使用本地 JSON 或写死数据。",
            "通知推送不需要系统级推送，只要在界面里展示分类结果。",
        ]
    )

    add_h1(doc, "4. 三人协作分工")
    add_fixed_table(
        doc,
        [
            ("A：前端界面", "时间轴、任务卡片、状态面板、通知防火墙、重排动画、移动端适配。"),
            ("B：Agent 决策逻辑", "规则引擎或 API 调用；输入任务、状态、外部事件，输出新排期和解释。"),
            ("C：Demo 与路演", "模拟数据、演示脚本、3 分钟讲稿、评委可能问题、现场彩排控时。"),
        ],
        headers=["成员", "负责内容"],
        widths=[Inches(1.6), Inches(4.85)],
        header_fill="E8EEF5",
    )

    add_h1(doc, "5. 明日执行时间表")
    add_fixed_table(
        doc,
        [
            ("10:00-10:30", "定最终故事", "只确定一个主线：状态变差 + 紧急任务插入 + Agent 智能重排。"),
            ("10:30-13:00", "搭主界面", "完成时间轴、状态卡片、任务卡片、基础排期展示。"),
            ("13:00-15:00", "做重排闭环", "一键触发突发事件，自动展示新排期和通知分流。"),
            ("15:00-16:30", "接 Agent 解释", "把重排理由、对话入口、晨报摘要跑通。"),
            ("16:30-17:30", "视觉与数据", "补充状态标签、模拟任务、演示细节和动效。"),
            ("17:30-18:30", "路演彩排", "按 3 分钟版本反复演示，修掉卡顿点。"),
            ("18:30-19:00", "冻结功能", "只修 Bug，不加新功能。"),
        ],
        headers=["时间", "目标", "产出"],
        widths=[Inches(1.15), Inches(1.3), Inches(4.0)],
        header_fill="E8EEF5",
    )

    add_h1(doc, "6. 推荐 Demo 剧本")
    add_numbered(
        doc,
        [
            "展示今日节律首页：上午是专注时段，下午有任务，晚上是低打扰休息时段。",
            "点击“昨晚睡眠差”：系统自动降低上午任务强度，插入恢复窗口。",
            "点击“GitHub 插入 P1 Issue”：系统判断任务紧急程度和当前状态。",
            "Agent 给出两个方案：现在 25 分钟最小修复，或推迟到 14:00 并后移低优任务。",
            "展示注意力防火墙：紧急事项立即提示，普通消息延后，夜间消息进入静默归档。",
            "切到晨报：展示夜间摘要、今日排期、睡眠状态导致的调整原因。",
        ]
    )
    add_callout(
        doc,
        "演示冲击点",
        "评委需要看到一个清晰瞬间：用户状态变差 + 外部紧急任务插入时，Agent 没有粗暴提醒，而是判断打扰等级、自动重排，并解释原因。",
        fill="FFF8E7",
    )

    add_h1(doc, "7. 评委眼前一亮的亮点表达")
    add_fixed_table(
        doc,
        [
            ("注意力防火墙", "不是提醒器，而是打扰管理器。系统决定什么时候提醒、什么时候延迟、什么时候静默。"),
            ("状态感知排期", "不是按空闲时间排任务，而是按用户的真实精力排任务。"),
            ("夜间静默晨报", "夜间 24h 后台吸收外部变化，但不打扰用户，早晨一次交付摘要。"),
            ("可解释排期", "每次重排都有原因，降低用户对 Agent 黑箱决策的不信任。"),
            ("自主权旋钮", "普通调整自动执行，重大调整先询问，睡眠和课程等锚点不可侵犯。"),
            ("个人节律画像", "系统越用越懂用户什么时候最适合学习、沟通、复盘或休息。"),
        ],
        headers=["亮点", "怎么讲"],
        widths=[Inches(1.55), Inches(4.9)],
        header_fill="E8EEF5",
    )

    add_h1(doc, "8. 路演话术")
    add_para(
        doc,
        "开场：现在很多 AI 日历都在帮人把时间排满，但真实用户的问题不是没有提醒，而是被提醒、消息和任务切碎。我们的项目想解决的是：AI 不仅要知道该做什么，还要知道什么时候不该打扰你。",
        bold_prefix="开场：",
    )
    add_para(
        doc,
        "方案：我们做的是一个 24 小时状态感知的时间调度 Agent。它读取课表、任务、外部平台变化和用户状态，动态判断现在该做什么、是否要重排、是否应该通知用户。",
        bold_prefix="方案：",
    )
    add_para(
        doc,
        "差异：它不是普通 AI 日历，而是注意力防火墙 + 状态感知排期 + 夜间静默晨报。白天主动推进，休息时低打扰，夜间静默吸收，晨起一次交付。",
        bold_prefix="差异：",
    )
    add_para(
        doc,
        "收束：我们希望这个 Agent 不只是管理任务，而是学习人的节奏，让用户在该专注时更专注，该休息时真的能休息。",
        bold_prefix="收束：",
    )

    add_h1(doc, "9. 最终建议")
    add_callout(
        doc,
        "明天的产品策略",
        "坚决砍功能，保留一个强故事闭环。把完整产品愿景留给路演，把现场开发集中在“状态变化触发智能重排”这个核心瞬间。",
        fill="F0FDF4",
    )
    add_bullets(
        doc,
        [
            "主标题建议：24h Personal Rhythm Agent。",
            "中文副标题：一个知道什么时候不该打扰你的时间调度智能体。",
            "现场展示优先级：界面完整 > 重排闭环 > 解释清楚 > 真实接入。",
            "失败风险最高的部分是范围过大，必须避免临场加功能。",
        ]
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
